from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DOCXRenderer:


    def render(
        self,
        paper_data,
        output_path
    ):


        document = Document()


        # HEADER

        header = document.sections[0].header

        paragraph = header.paragraphs[0]

        paragraph.text = "QuestAI Generated Question Paper"

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


        paragraph.runs[0].font.size = Pt(14)



        # BODY


        for line in paper_data:


            p = document.add_paragraph()


            p.add_run(
                str(line)
            ).font.size = Pt(12)



        # FOOTER

        footer = document.sections[0].footer


        footer.paragraphs[0].text = (
            "Generated using QuestAI"
        )


        document.save(
            output_path
        )


        return output_path