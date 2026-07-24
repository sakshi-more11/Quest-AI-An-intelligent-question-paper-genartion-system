from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT



class TemplateDOCXRenderer:


    def render(
        self,
        template,
        metadata,
        questions,
        output_path
    ):


        document = Document()



        # -------------------------
        # PAGE SETTINGS
        # -------------------------

        section = document.sections[0]

        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)



        # -------------------------
        # HEADER
        # -------------------------


        header = document.add_paragraph()


        header.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )


        run = header.add_run(

            metadata["college_name"]

        )


        run.bold=True

        run.font.size = Pt(14)



        exam = document.add_paragraph()


        exam.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )


        exam.add_run(

            "End Semester Examination"

        ).bold=True




        # -------------------------
        # METADATA TABLE
        # -------------------------


        table = document.add_table(

            rows=2,

            cols=2

        )


        table.alignment = (
            WD_TABLE_ALIGNMENT.CENTER
        )


        table.cell(0,0).text = (
            "Course Code"
        )


        table.cell(0,1).text = (

            metadata["course_code"]

        )


        table.cell(1,0).text = (

            "Course Name"

        )


        table.cell(1,1).text = (

            metadata["course_name"]

        )



        document.add_paragraph()



        # -------------------------
        # QUESTION TABLE
        # -------------------------


        q_table = document.add_table(

            rows=1,

            cols=5

        )


        q_table.alignment = (
            WD_TABLE_ALIGNMENT.CENTER
        )



        headers=[

            "Question",

            "Description",

            "Marks",

            "CO",

            "BL"

        ]


        for i,h in enumerate(headers):

            q_table.cell(
                0,i
            ).text=h




        for q in questions:


            row=q_table.add_row().cells



            row[0].text = (
                q["number"]
            )


            row[1].text = (
                q["text"]
            )


            row[2].text = (
                str(q["marks"])
            )


            row[3].text = (
                q["co"]
            )


            row[4].text = (
                q["bl"]
            )




        # -------------------------
        # FOOTER
        # -------------------------


        footer = (
            document.sections[0]
            .footer
        )


        footer.paragraphs[0].text = (

            "Signature                         Page Number"

        )


        document.save(
            output_path
        )



        return output_path