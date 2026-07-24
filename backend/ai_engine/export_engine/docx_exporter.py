from docx import Document
import os


class DOCXExporter:


    def export(self, paper, output_path):

        os.makedirs(
            os.path.dirname(output_path),
            exist_ok=True
        )


        doc = Document()


        doc.add_heading(
            paper["title"],
            level=1
        )


        doc.add_paragraph(
            f"Subject : {paper['subject']}"
        )


        doc.add_paragraph(
            f"Duration : {paper['duration']}"
        )


        doc.add_paragraph(
            f"Total Marks : {paper['total_marks']}"
        )


        for section in paper["sections"]:

            doc.add_heading(
                section["name"],
                level=2
            )


            for index,question in enumerate(
                section["questions"],
                start=1
            ):

                doc.add_paragraph(
                    f"{index}. {question['question']} "
                    f"({question['marks']} Marks)"
                )


        doc.save(output_path)


        return output_path