from backend.ai_engine.template_learning.template_parser import TemplateParser
from backend.ai_engine.template_learning.pattern_extractor import PatternExtractor
from backend.ai_engine.template_learning.template_analyzer import TemplateAnalyzer
from pathlib import Path


class TemplateLearningPipeline:

    def __init__(self):

        self.parser = TemplateParser()

        self.extractor = PatternExtractor()

        self.analyzer = TemplateAnalyzer()

    def learn(self, file_path):
        print("<<<<<<<<<<<< TEMPLATE PIPELINE RUNNING >>>>>>>>>>>>>>")

        print("\n========== STEP 1 ==========")
        print("Parsing Previous Year Paper...")

        parsed = self.parser.parse(file_path)
        text = parsed["text"]

        print("\n========== STEP 2 ==========")
        print("Extracting Paper Pattern...")

        questions = self.extractor.extract(text)

        print("\n========== STEP 3 ==========")
        print("Analyzing Template...")

        template = self.analyzer.analyze(questions)

        # Preserve the parser's slot details (including OR pairing) while the
        # analyser contributes aggregate metrics for the UI.
        template.update(parsed["learned_template"])

        print("\nTemplate Learning Completed!")

        template["source_path"] = str(file_path)
        template["source_type"] = Path(file_path).suffix.lower().lstrip(".")
        template["layout_profile"] = self._layout_profile(file_path)
        return template

    @staticmethod
    def _layout_profile(file_path):
        """Capture reusable native-DOCX layout information without flattening it."""
        path = Path(file_path)
        if path.suffix.lower() != ".docx":
            return {"fidelity_mode": "pdf-reference", "source": str(path)}
        try:
            from docx import Document
            document = Document(path)
            section = document.sections[0]
            return {
                "fidelity_mode": "docx-clone",
                "page": {"width": section.page_width, "height": section.page_height,
                         "top_margin": section.top_margin, "bottom_margin": section.bottom_margin,
                         "left_margin": section.left_margin, "right_margin": section.right_margin},
                "styles": {style.name: {"font": style.font.name, "size": style.font.size.pt if style.font.size else None}
                           for style in document.styles if style.type == 1},
                "header_text": [paragraph.text for paragraph in section.header.paragraphs],
                "footer_text": [paragraph.text for paragraph in section.footer.paragraphs],
                "tables": [{"rows": len(table.rows), "columns": len(table.columns)} for table in document.tables],
            }
        except Exception as exc:
            return {"fidelity_mode": "source-preserved", "source": str(path), "analysis_error": str(exc)}
