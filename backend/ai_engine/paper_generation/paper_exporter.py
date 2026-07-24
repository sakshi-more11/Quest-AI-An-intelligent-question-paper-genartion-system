"""
paper_exporter.py
"""

import os
import json
from pathlib import Path

from docx import Document

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth


class PaperExporter:

    def __init__(self):

        self.export_dir = "exports"

        os.makedirs(

            self.export_dir,

            exist_ok=True

        )

    # -------------------------------------

    # JSON

    # -------------------------------------

    def export_json(self, paper):

        path = os.path.join(

            self.export_dir,

            "question_paper.json"

        )

        with open(

            path,

            "w",

            encoding="utf-8"

        ) as f:

            json.dump(

                paper,

                f,

                indent=4,

                ensure_ascii=False

            )

        return path

    # -------------------------------------

    # DOCX

    # -------------------------------------

    def export_docx(self, paper, template_path=None):

        path = os.path.join(

            self.export_dir,

            "question_paper.docx"

        )

        template_path = template_path or paper.get("template_path")
        doc = Document(template_path) if template_path and Path(template_path).suffix.lower() == ".docx" else Document()

        if template_path and Path(template_path).suffix.lower() == ".docx":
            has_question_placeholder = self._has_question_placeholder(doc)
            self._replace_template_fields(doc, paper)
            if not has_question_placeholder:
                self._append_questions(doc, paper)
            doc.save(path)
            return path

        doc.add_heading(

            paper["title"],

            level=1

        )

        doc.add_paragraph(

            f"Subject : {paper['subject']}"

        )

        doc.add_paragraph(

            f"Duration : {paper['duration']}"

        )

        doc.add_paragraph(

            f"Total Marks : {paper['total_marks']}"

        )

        doc.add_paragraph()

        for section in paper["sections"]:

            doc.add_heading(

                section["name"],

                level=2

            )

            for i, q in enumerate(

                section["questions"],

                start=1

            ):

                doc.add_paragraph(

                    f"{i}. {q['question']} ({q['marks']} Marks)"

                )

        doc.save(path)

        return path

    # -------------------------------------

    # PDF

    # -------------------------------------

    def export_pdf(self, paper):

        path = os.path.join(

            self.export_dir,

            "question_paper.pdf"

        )

        pdf = canvas.Canvas(path, pagesize=A4)
        page_width, page_height = A4
        left, right, bottom = 50, page_width - 50, 55
        y = page_height - 50

        def new_page():
            pdf.showPage()
            return page_height - 50

        def draw_wrapped(text, x, font, size, leading=15):
            nonlocal y
            pdf.setFont(font, size)
            words, line = str(text).split(), ""
            lines = []
            for word in words:
                candidate = (line + " " + word).strip()
                if line and stringWidth(candidate, font, size) > right - x:
                    lines.append(line)
                    line = word
                else:
                    line = candidate
            lines.append(line or " ")
            for line in lines:
                if y < bottom + leading:
                    y = new_page()
                    pdf.setFont(font, size)
                pdf.drawString(x, y, line)
                y -= leading

        pdf.setFont(

            "Helvetica-Bold",

            16

        )

        pdf.drawString(

            left,

            y,

            paper["title"]

        )

        y -= 30

        pdf.setFont(

            "Helvetica",

            12

        )

        pdf.drawString(

            left,

            y,

            f"Subject : {paper['subject']}"

        )

        y -= 20

        pdf.drawString(

            left,

            y,

            f"Duration : {paper['duration']}"

        )

        y -= 20

        pdf.drawString(

            left,

            y,

            f"Total Marks : {paper['total_marks']}"

        )

        y -= 40

        for section in paper["sections"]:

            pdf.setFont(

                "Helvetica-Bold",

                14

            )

            if y < bottom + 30:
                y = new_page()
            pdf.drawString(left, y, section["name"])

            y -= 25

            pdf.setFont(

                "Helvetica",

                11

            )

            for i, q in enumerate(

                section["questions"],

                start=1

            ):

                if q.get("or_before"):
                    if y < bottom + 20:
                        y = new_page()
                    pdf.setFont("Helvetica-Bold", 11)
                    pdf.drawCentredString(page_width / 2, y, "OR")
                    y -= 18
                number = str(q.get("question_no") or i)
                sub = str(q.get("sub_question") or "").strip()
                label = f"{number} ({sub})" if sub else number
                draw_wrapped(f"{label}. {q.get('question', '')} ({q.get('marks', '')} Marks)", 60, "Helvetica", 11)

        pdf.save()

        return path

    @staticmethod
    def _question_text(paper):
        lines = []
        for section in paper.get("sections", []):
            lines.append(section.get("name", "Section"))
            for index, question in enumerate(section.get("questions", []), 1):
                if question.get("or_before"):
                    lines.append("OR")
                lines.append(f"{index}. {question.get('question', '')} ({question.get('marks', '')} Marks)")
        return "\n".join(lines)

    def _replace_template_fields(self, doc, paper):
        replacements = {
            "{{TITLE}}": paper.get("title", "University Examination"),
            "{{SUBJECT}}": paper.get("subject", ""),
            "{{DURATION}}": paper.get("duration", ""),
            "{{TOTAL_MARKS}}": str(paper.get("total_marks", "")),
            "{{QUESTIONS}}": self._question_text(paper),
        }
        containers = list(doc.paragraphs)
        for table in doc.tables:
            containers.extend(cell.paragraphs for row in table.rows for cell in row.cells)
        for section in doc.sections:
            containers.extend(section.header.paragraphs)
            containers.extend(section.footer.paragraphs)
        for paragraph in containers:
            text = paragraph.text
            if not any(token in text for token in replacements):
                continue
            for token, value in replacements.items():
                text = text.replace(token, value)
            # Retain the paragraph's existing style and the first run's typography.
            if paragraph.runs:
                paragraph.runs[0].text = text
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(text)

    @staticmethod
    def _has_question_placeholder(doc):
        return any("{{QUESTIONS}}" in paragraph.text for paragraph in doc.paragraphs)

    def _append_questions(self, doc, paper):
        doc.add_paragraph()
        for section in paper.get("sections", []):
            # Faculty DOCX files frequently do not contain Word's built-in
            # Heading styles.  Adding one then raises KeyError and aborts the
            # entire paper API.  Use a normal paragraph and preserve the
            # template's available typography instead.
            heading = doc.add_paragraph()
            heading.add_run(section.get("name", "Section")).bold = True
            for index, question in enumerate(section.get("questions", []), 1):
                if question.get("or_before"):
                    doc.add_paragraph("OR").alignment = 1
                number = str(question.get("question_no") or index)
                sub = str(question.get("sub_question") or "").strip()
                label = f"{number} ({sub})" if sub else number
                doc.add_paragraph(f"{label}. {question.get('question', '')} ({question.get('marks', '')} Marks)")
